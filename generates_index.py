#Importing modules
import mysql.connector
import docx
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches

# Indexes output path
output_path = "C:/Users/gdatn/Documentos/python/estudos/Atividades/pdfProject/indexes"

# Function to generate indexes 
def generates_index(pdf_id, name) :

    query = "SELECT word FROM special_word WHERE id_pdf_information = %d" %pdf_id
    print(query)
    cursor.execute(query)
    special_words = cursor.fetchall()
    print(special_words)

    query = "SELECT number_top_words FROM pdf_information WHERE id = %d" %pdf_id
    print(query)
    cursor.execute(query)
    lines = cursor.fetchone()
    number = lines[0]

    query = "SELECT word, count FROM word_count WHERE id_pdf_information = %d ORDER BY count DESC LIMIT %d" %(pdf_id, number)
    print(query)
    cursor.execute(query)
    top_words_count = cursor.fetchall()
    top_words = []
    for words in top_words_count :
        top_words.append(words[0])

    document = Document()

    styles = document.styles

    # Style Paragraph 1
    p = styles.add_style("Paragraph1", WD_STYLE_TYPE.PARAGRAPH)
    p.font.name = "Courier New"
    p.font.size = Pt(11)
    p.font.color.rgb = RGBColor(0, 0, 255)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    # Style Paragraph 2
    q = styles.add_style("Paragraph2", WD_STYLE_TYPE.PARAGRAPH)
    q.font.name = "Courier New"
    q.font.size = Pt(11)
    q.font.color.rgb = RGBColor(139, 0, 0)
    q.paragraph_format.space_before = Pt(0)
    q.paragraph_format.space_after = Pt(0)
    q.paragraph_format.line_spacing = 1
    # Style Heading 1
    h1 = styles.add_style("H1", WD_STYLE_TYPE.PARAGRAPH)
    h1.base_style = styles["Heading 1"]
    h1.font.name = "Courier New"
    h1.font.size = Pt(12)
    h1.font.bold = False
    h1.font.color.rgb = RGBColor(0, 0, 128)
    h1.paragraph_format.space_before = Pt(10)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.line_spacing = 1

    # Add Title
    document.add_heading("PDF: " + name, 0)
    document.add_heading("Index of main words")

    # Add paragraph with H1
    document.add_paragraph("Special Words: ", h1)

    for word_array in special_words :
        # Add paragraph 2 - font color = red
        word = word_array[0]
        word = word.strip()
        query = "SELECT DISTINCT word, page_number FROM word_page_number WHERE id_pdf_information = %d AND word = '%s'" %(pdf_id, word)
        cursor.execute(query)
        print(query)
        page_numbers = cursor.fetchall()
        pages = ""
        for page in page_numbers :
            if page == page_numbers[-1]:
                pages = pages + str(page[1]) + "."
            else :
                pages = pages + str(page[1]) + ", "
        space_word = 30 - len(word)
        document.add_paragraph(word+"."*space_word+" page: " +pages, style="Paragraph2")

    # Add paragraph with H1
    document.add_paragraph("Top Words: ", style="H1")

    for word in top_words :
        # Add paragraph 1 - font color = blue
        query = "SELECT DISTINCT word, page_number FROM word_page_number WHERE id_pdf_information = %d AND word = '%s'" %(pdf_id, word)
        print(query)
        cursor.execute(query)
        page_numbers = cursor.fetchall()
        pages = ""
        for page in page_numbers :
            if page == page_numbers[-1]:
                pages = pages + str(page[1]) + "."
            else :
                pages = pages + str(page[1]) + ", "
        space_word = 30 - len(word)
        document.add_paragraph(word+"."*space_word+" page: "+pages, style="Paragraph1")

    docx_archive = output_path + "/indexOf_"+str(name)+".docx"
    document.save(docx_archive)

    return()

#Connecting to MySQL

HOST = input("Enter MySQL Host: ")
USER = input("Enter your MySQL User: ")
PASSWD = input("Enter the password: ")

global con
con = mysql.connector.connect(host = "localhost", user = "root", passwd = "12345")

control = True

while control :

    #If connection is on, execute following
    if con.is_connected() :

        #Printing info about MySQL Server
        db_info = con.get_server_info()
        print("\nConnected in MySQL v", db_info,"\n")
        
        #Establishing cursor
        cursor = con.cursor()

        query = "USE pdf_project"
        cursor.execute(query)

        answer = input("\nDo you want to generate indexes for all registered pdfs? Enter [Y] if yes or [N] to choose one: ")

        if answer == "Y" :
            query = "SELECT id, name FROM pdf_information"
            cursor.execute(query)
            info = cursor.fetchall()

            for row in info :
                pdf_id = row[0]
                name = row[1]
                generates_index(pdf_id, name)
            
            print("\nIndexes generated successfully! Open the Indexes folder to view. ")

        else :
            query = "SELECT id, name FROM pdf_information"
            cursor.execute(query)
            info = cursor.fetchall()

            print("\nChoose one PDF below to generate the index\n")

            for row in info :
                print(str(row[0]) + " - " + str(row[1]))

            pdf_id = int(input("\nEnter the related PDF number shown above:"))

            # Função para tomar informação a partir da variável info (função index)
            query = "SELECT name FROM pdf_information WHERE id = %d" %pdf_id
            cursor.execute(query)
            name = cursor.fetchone()[0]
            print("Antes de chamar a função nome PDF = ", name)

            generates_index(pdf_id, name)
            print("\nIndex generated successfully! Open the Indexes folder to view.\n")

    answer = input("\nDo you want to do some more analysis? Enter [Y] if yes or [N] to exit the program: ")

    if answer != "Y" : 
        control = False

#Commiting changes
con.commit()

#If the connection is still opened, close the connection
if con.is_connected() :
    cursor.close()
    con.close()
    print("\nMySQL connection closed.")